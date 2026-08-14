# -*- coding: utf-8 -*-
"""混合估值评分：部分行业用PB、其余用PE"""
import pandas as pd, numpy as np, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

out_dir = r'D:\CC\Mid\估值\申万'
sw1_dir = r'D:\CC\Mid\估值\中信\SW1'
C_PRI=RGBColor(27,54,93);C_SEC=RGBColor(74,119,122);C_RED=RGBColor(192,57,43);C_GRN=RGBColor(39,174,96)

PB_INDS = {'银行','非银金融','电力设备','交通运输','有色金属','建筑装饰','食品饮料','家用电器','商贸零售'}

# ── 1. 读取数据 ──
yr = pd.read_excel(os.path.join(sw1_dir, '申万一级行业年度收益率矩阵.xlsx'), index_col=0)
qr = pd.read_excel(os.path.join(sw1_dir, '申万一级行业季度收益率矩阵.xlsx'), index_col=0)
# 分离行业名称和代码（名称用于merge，代码留存备用）
yri = yr.index.astype(str)
yr.index = yri.str.split('\n').str[0].str.strip()
yr_code = yri.str.split('\n').str[1].str.strip()  # 行业代码
qri = qr.index.astype(str)
qr.index = qri.str.split('\n').str[0].str.strip()
qr_code = qri.str.split('\n').str[1].str.strip()  # 行业代码
pe_stats = pd.read_excel(os.path.join(out_dir, '申万行业PE统计.xlsx'))
pb_stats = pd.read_excel(os.path.join(out_dir, '申万行业PB统计.xlsx'))

# ── 2. 收益率统计 ──
ny = yr.shape[1]
stats = pd.DataFrame({
    '年化平均收益率(%)': yr.mean(1).round(2),
    '年化波动率(%)': yr.std(1).round(2),
    '最大年涨幅(%)': yr.max(1).round(2),
    '最大年跌幅(%)': yr.min(1).round(2),
    '累计收益率(%)': ((yr/100+1).prod(1)-1).mul(100).round(2)})
stats['年胜率(%)'] = (yr > 0).sum(1) / ny * 100
stats['收益风险比'] = (stats['累计收益率(%)'] / (-stats['最大年跌幅(%)'].clip(upper=-0.01))).round(2)

latest_yr_cols = [c for c in yr.columns if c >= '2024']
stats['近3年累计收益(%)'] = ((yr[latest_yr_cols]/100+1).prod(1)-1).mul(100).round(2)
if '2026' in yr.columns: stats['2026年收益(%)'] = yr['2026'].round(2)
if '2025' in yr.columns: stats['2025年收益(%)'] = yr['2025'].round(2)
latest_q = qr.columns[-1] if len(qr.columns) > 0 else ''
if latest_q: stats['最新季度收益(%)'] = qr[latest_q].round(2)
if len(qr.columns) >= 4:
    stats['近4季累计收益(%)'] = ((qr[qr.columns[-4:]]/100+1).prod(1)-1).mul(100).round(2)

stats = stats.reset_index().rename(columns={'index': '行业'})

# ── 3. 合并PE/PB ──
merged = stats.merge(pe_stats, on='行业', how='left').merge(pb_stats, on='行业', how='left')

# ── 4. 评分 ──
merged['PE评分'] = pd.cut(merged['PE历史分位数(%)'].fillna(50), bins=[-1,30,50,70,100], labels=[3,2,1,0]).astype(int)
merged['PB评分'] = pd.cut(merged['PB历史分位数(%)'].fillna(50), bins=[-1,30,50,70,100], labels=[3,2,1,0]).astype(int)

mom_cols = [c for c in ['近3年累计收益(%)','2026年收益(%)','最新季度收益(%)'] if c in merged.columns]
if mom_cols:
    merged['动量综合'] = merged[mom_cols].mean(1)
    merged['动量评分'] = pd.qcut(merged['动量综合'].rank(method='first'), q=4, labels=[0,1,2,3]).astype(int)
else:
    merged['动量评分'] = 1; merged['动量综合'] = 0

merged['风险评分'] = (merged['年化波动率(%)'].rank(pct=True) * 3).round(0).clip(0, 3).astype(int)

# ── 混合估值评分 ──
def hybrid_val(row):
    ind = row['行业']
    if ind in PB_INDS:
        return row['PB评分'], 'PB'
    else:
        return row['PE评分'], 'PE'

merged[['估值评分','估值依据']] = merged.apply(lambda r: pd.Series(hybrid_val(r)), axis=1)
merged['综合评分'] = (merged['估值评分']*0.4 + merged['动量评分']*0.4 + (3-merged['风险评分'])*0.2).round(2)

def classify(s):
    if s >= 2.0: return '积极配置'
    elif s >= 1.5: return '标配'
    elif s >= 1.0: return '谨慎观望'
    else: return '规避'
merged['投资评级'] = merged['综合评分'].apply(classify)

def pct_label(p): return '深度低估' if p<=20 else ('低估' if p<=40 else ('合理' if p<=60 else ('偏高' if p<=80 else '高估')))
merged['PE状态'] = merged['PE历史分位数(%)'].fillna(50).apply(pct_label)
merged['PB状态'] = merged['PB历史分位数(%)'].fillna(50).apply(pct_label)
merged['动量状态'] = merged['动量评分'].apply(lambda r: {3:'强势',2:'偏强',1:'偏弱',0:'弱势'}.get(r,'未知'))
merged['风险等级'] = merged['风险评分'].apply(lambda r: {0:'低',1:'低',2:'中',3:'高'}.get(r,'未知'))
merged['估值状态'] = merged.apply(lambda r: r['PB状态'] if r['估值依据']=='PB' else r['PE状态'], axis=1)

merged = merged.sort_values('综合评分', ascending=False)

# ── 5. 导出Excel ──
ex_cols = ['行业','投资评级','综合评分','估值评分','动量评分','风险等级',
           '估值状态','估值依据',
           '最新PE','PE历史分位数(%)','PE状态','最新PB','PB历史分位数(%)','PB状态',
           '动量状态','2026年收益(%)','近3年累计收益(%)','年化波动率(%)','年胜率(%)','收益风险比']
ex_cols = [c for c in ex_cols if c in merged.columns]
from openpyxl.styles import Font, Border, Side, Alignment, numbers, PatternFill
from openpyxl.utils import get_column_letter

excel_path = os.path.join(out_dir, '行业投资价值综合评估_混合版.xlsx')
with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
    merged[ex_cols].to_excel(writer, sheet_name='混合评估', index=False)
    ws = writer.sheets['混合评估']
    thin = Side(style='thin', color='999999')
    border = Border(top=thin, left=thin, right=thin, bottom=thin)
    font_msyh = Font(name='Microsoft YaHei', size=10)
    font_msyh_bold = Font(name='Microsoft YaHei', size=10, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1B365D', end_color='1B365D', fill_type='solid')
    align_center = Alignment(horizontal='center', vertical='center')

    # 格式化表头
    for cell in ws[1]:
        cell.font = font_msyh_bold
        cell.fill = header_fill
        cell.border = border
        cell.alignment = align_center

    # 格式化数据行
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.font = font_msyh
            cell.border = border
            cell.alignment = align_center
            # 数值保留两位小数
            if isinstance(cell.value, (int, float)):
                cell.number_format = '0.00'

    # 列宽
    for i, col in enumerate(ex_cols, 1):
        col_letter = get_column_letter(i)
        max_len = max(len(str(col)), merged[col].astype(str).str.len().max() if merged[col].dtype=='object' else 10)
        ws.column_dimensions[col_letter].width = min(max_len+4, 30)
print(f"导出: {excel_path}")

# ════════════════════════════════════════════════════════════
# 6. Word报告
# ════════════════════════════════════════════════════════════
def sf(run, name='Microsoft YaHei', size=Pt(10.5), color=RGBColor(51,51,51), bold=False, italic=False):
    run.font.name = name; run.font.size = size; run.font.color.rgb = color; run.bold = bold; run.italic = italic
    rPr = run._r.get_or_add_rPr(); rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}"/>'))

def add_title(doc, text, size=Pt(22)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); sf(r, size=size, color=C_PRI, bold=True)

def add_h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f'一、{text}'); sf(r, size=Pt(14), color=C_PRI, bold=True)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="1B365D"/></w:pBdr>'))

def add_h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'▎{text}'); sf(r, size=Pt(11), color=C_SEC, bold=True)

def add_body(doc, text, bf=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.7)
    if bf: r = p.add_run(bf); sf(r, bold=True, color=C_PRI)
    r = p.add_run(text); sf(r)

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def fmt_table(t):
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(8)
            if i == 0:
                shade(cell, '1B365D')
                for p in cell.paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(255,255,255); r.bold = True

def add_table(doc, df, title, cols):
    add_h2(doc, title)
    n = len(df) + 1; t = doc.add_table(rows=n, cols=len(cols)); t.style = 'Table Grid'
    for j, h in enumerate(cols): t.cell(0, j).text = h
    for i, (_, row) in enumerate(df.iterrows()):
        for j, c in enumerate(cols):
            v = row.get(c, '')
            t.cell(i+1, j).text = str(v) if not pd.isna(v) else '-'
            if c == '投资评级' and isinstance(v, str):
                clr = C_GRN if '积极' in v else (C_SEC if '标配' in v else (RGBColor(200,150,30) if '观望' in v else C_RED))
                for p in t.cell(i+1, j).paragraphs:
                    for r in p.runs: r.font.color.rgb = clr; r.bold = True
    fmt_table(t)

doc = Document()
add_title(doc, '申万一级行业综合投资价值与风险评估报告（混合版）')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('银行/非银金融/电力设备/交通运输/有色金属/建筑装饰/食品饮料/家用电器/商贸零售 → PB估值，其余 → PE估值'); sf(r, size=Pt(10), color=C_SEC)
doc.add_paragraph()

add_h1(doc, '分析框架说明')
add_body(doc, '9个行业使用PB历史百分位评分（银行、非银金融、电力设备、交通运输、有色金属、建筑装饰、食品饮料、家用电器、商贸零售）；其余22个行业使用PE历史百分位评分。')
add_body(doc, '综合评分=估值评分×40%+动量评分×40%+(3-风险评分)×20%，满分3分。')

add_h1(doc, '综合排名')
rc = ['行业','投资评级','综合评分','估值评分','动量评分','风险等级','估值状态','估值依据']
add_table(doc, merged.head(31), '全行业排名', rc)

for level, label in [('积极配置','积极配置板块'),('标配','标配板块'),('谨慎观望','谨慎观望板块'),('规避','规避板块')]:
    sub = merged[merged['投资评级']==level]
    if len(sub): add_h1(doc, label); add_table(doc, sub[rc], '', rc)

add_h1(doc, '分维度分析')
val_disp = merged[['行业','估值状态','估值依据']].copy()
val_disp['估值百分位'] = merged.apply(lambda r: f"{r['PE历史分位数(%)']:.1f}%" if r['估值依据']=='PE' else f"{r['PB历史分位数(%)']:.1f}%", axis=1)
val_disp['最新值'] = merged.apply(lambda r: r['最新PE'] if r['估值依据']=='PE' else r['最新PB'], axis=1)
add_table(doc, val_disp.sort_values('估值百分位'), '估值由低到高', ['行业','估值状态','估值百分位','最新值','估值依据'])

mom_d = [c for c in ['行业','动量状态','2026年收益(%)','近3年累计收益(%)','最新季度收益(%)'] if c in merged.columns]
add_table(doc, merged.sort_values('动量综合', ascending=False)[mom_d], '动量排名', mom_d)
add_table(doc, merged.sort_values('年化波动率(%)')[['行业','风险等级','年化波动率(%)','最大年跌幅(%)','年胜率(%)','收益风险比']], '风险排名', ['行业','风险等级','年化波动率(%)','最大年跌幅(%)','年胜率(%)','收益风险比'])

add_h1(doc, '估值依据说明')
add_body(doc, '银行/非银金融：高杠杆行业，净资产价值为核心定价基础，PB比PE更稳定。')
add_body(doc, '电力设备/交通运输/建筑装饰：典型重资产行业，资产价值驱动，PB参考意义更大。')
add_body(doc, '有色金属：强周期行业，盈利波动剧烈，PE易失真，PB更可靠。')
add_body(doc, '食品饮料/家用电器：估值已进入成熟期，PB作为辅助参考更全面。')
add_body(doc, '商贸零售：行业面临转型，盈利波动大，PB更能反映资产底牌。')

add_h1(doc, '投资策略总结')
active = merged[merged['投资评级']=='积极配置']['行业'].tolist()
neutral = merged[merged['投资评级']=='标配']['行业'].tolist()
av = merged[merged['投资评级'].isin(['谨慎观望','规避'])]['行业'].tolist()
if active: add_body(doc, '、'.join(active)+'。', '积极配置（超配）：')
if neutral: add_body(doc, '、'.join(neutral)+'。', '标配：')
if av: add_body(doc, '、'.join(av)+'。', '谨慎/规避（低配）：')

out = os.path.join(out_dir, '申万一级行业综合投资价值与风险评估报告_混合版.docx')
doc.save(out); print(f"导出: {out}")
print("\n全部完成！")
