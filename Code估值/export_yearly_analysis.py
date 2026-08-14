# -*- coding: utf-8 -*-
"""导出年度收益率分析报告 Word 文档"""
import pandas as pd, numpy as np, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

out_dir = r'D:\CC\Mid\估值\申万'
df = pd.read_excel(os.path.join(out_dir,'申万一级行业年度收益率矩阵.xlsx'),index_col=0)

ny = df.shape[1]
s=pd.DataFrame({'平均收益率(%)':df.mean(1).round(2),'波动率(%)':df.std(1).round(2),
    '最大涨幅(%)':df.max(1).round(2),'最大跌幅(%)':df.min(1).round(2),
    '累计收益率(%)':((df/100+1).prod(1)-1).mul(100).round(2)})
s['胜率(%)']=(df>0).sum(1)/ny*100
s['收益风险比']=(s['累计收益率(%)']/(-s['最大跌幅(%)'].clip(upper=-0.01))).round(2)
s=s.sort_values('累计收益率(%)',ascending=False)

doc=Document()
C_PRI=RGBColor(27,54,93);C_SEC=RGBColor(74,119,122);C_RED=RGBColor(192,57,43);C_GRN=RGBColor(39,174,96)

def sf(run,name='Microsoft YaHei',size=Pt(10.5),color=RGBColor(51,51,51),bold=False,italic=False):
    run.font.name=name;run.font.size=size;run.font.color.rgb=color;run.bold=bold;run.italic=italic
    r=run._r.get_or_add_rPr();r.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}"/>'))

def add_t(text,size=Pt(22)):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text);sf(r,size=size,color=C_PRI,bold=True)

def h1(text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(16);p.paragraph_format.space_after=Pt(8)
    r=p.add_run(f'一、{text}');sf(r,size=Pt(14),color=C_PRI,bold=True)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="1B365D"/></w:pBdr>'))

def h2(text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f'▎{text}');sf(r,size=Pt(11),color=C_SEC,bold=True)

def body(text,bf=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(4);p.paragraph_format.first_line_indent=Cm(0.7)
    if bf:r=p.add_run(bf);sf(r,bold=True,color=C_PRI)
    r=p.add_run(text);sf(r)

def shade(cell,color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def ft(t):
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate(t.rows):
        for j,cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:r.font.size=Pt(8.5)
            if i==0:
                shade(cell,'1B365D')
                for p in cell.paragraphs:
                    for r in p.runs:r.font.color.rgb=RGBColor(255,255,255);r.bold=True

def tbl(data,title):
    h2(title);n=len(data)+1;t=doc.add_table(rows=n,cols=6);t.style='Table Grid'
    hs=['排名','行业','累计收益率','平均年收益','波动率','胜率']
    for j,h in enumerate(hs):t.cell(0,j).text=h
    for i,(idx,row) in enumerate(data.iterrows()):
        t.cell(i+1,0).text=str(i+1);t.cell(i+1,1).text=str(idx)
        t.cell(i+1,2).text=f"{row['累计收益率(%)']:.1f}%"
        t.cell(i+1,3).text=f"{row['平均收益率(%)']:.2f}%"
        t.cell(i+1,4).text=f"{row['波动率(%)']:.1f}%"
        t.cell(i+1,5).text=f"{row['胜率(%)']:.0f}%"
        c=C_GRN if row['累计收益率(%)']>0 else C_RED
        for j in[2,3]:
            for p in t.cell(i+1,j).paragraphs:
                for r in p.runs:r.font.color.rgb=c;r.bold=True
    ft(t)

add_t('申万一级行业年度收益率投资价值与风险评估报告')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f'基于2016-2026共{ny}个自然年度收益率分析');sf(r,size=Pt(11),color=C_SEC)
doc.add_paragraph()

h1('行业综合排名')
body(f'以下基于2016至2026年共{ny}个年度数据，对31个申万一级行业进行收益率统计分析。排名按累计收益率降序排列。')
tbl(s,'全部31个行业排名')

h1('行业投资价值分类')
h2('🟢 第一梯队：进攻型（高收益高波动）')
top6=s.head(6);tbl(top6,'Top 6行业')
body(('食品饮料(+280.5%)以年化复利领跑，驱动来自2016-2020年消费升级长期逻辑下的戴维斯双击。'
      '家用电器(+182.5%)紧随其后，受益于地产后周期与全球化出海。'
      '社会服务、电力设备、电子、医药生物均体现了A股核心资产的长牛特征。'
      '注意：部分行业近年出现明显均值回归。'))

h2('🟡 第二梯队：稳健型（绝对收益+低波动）')
mid=s.iloc[6:15];tbl(mid,'第7-15名')
body(('银行、公用事业等低波动行业累计收益稳定，具备高股息防御属性。'
      '银行胜率达60%+，是典型的稳健底仓品种。'
      '建筑装饰、石油石化等周期行业波动可控，适合作为组合中的卫星配置。'))

h2('🔴 第三梯队：困境型（大幅亏损等待拐点）')
bot=s.tail(8);tbl(bot,'排名靠后行业')
body(('房地产(-56.0%)累计跌幅最深，行业逻辑面临根本性转变。'
      '传媒(-46.7%)、纺织服饰(-37.2%)长期表现低迷。'
      '钢铁、商贸零售等传统行业受经济结构转型冲击显著。'))

h1('风险评估矩阵')
h2('高波动行业（年波动率>35%）')
tbl(s[s['波动率(%)']>35].sort_values('波动率(%)',ascending=False),'高波动行业')
h2('高回撤行业（最大单年跌幅>30%）')
tbl(s[s['最大跌幅(%)']<-30].sort_values('最大跌幅(%)'),'高回撤行业')
h2('收益风险比排名')
tbl(s.sort_values('收益风险比',ascending=False),'性价比排名')
body(('食品饮料以17.2的收益风险比位居榜首，每承担1%回撤可获得17.2%年化收益。'
      '家用电器、社会服务紧随其后，是大周期维度下最具性价比的赛道。'))

h1('投资策略建议')
body('食品饮料、电力设备逢回调定投，电子和医药生物长期配置。','长期核心配置：')
body('银行和公用事业提供稳定股息+低波动，作为组合压舱石。','稳健底仓配置：')
body('煤炭、有色金属在通胀/复苏阶段可阶段性超配。','周期轮动配置：')
body('房地产行业逻辑未理顺，传媒和纺织服饰等待基本面拐点信号。','需要规避：')

h1('年度收益率热力图')
img=os.path.join(out_dir,'申万一级行业年度收益率热力图.png')
if os.path.exists(img):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img,width=Inches(5.5))
    p2=doc.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p2.add_run(f'图：申万一级行业年度收益率热力图 (2016-2026)');sf(r,size=Pt(9),color=C_SEC,italic=True)

out=os.path.join(out_dir,'申万一级行业年度收益率分析报告.docx')
doc.save(out);print(f"导出: {out}")
