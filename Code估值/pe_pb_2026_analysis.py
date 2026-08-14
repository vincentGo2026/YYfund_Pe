# -*- coding: utf-8 -*-
"""基于最新PE/PB数据 + 收益率的综合投资价值与风险评估"""
import pandas as pd, numpy as np, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

out_dir = r'D:\CC\Mid\估值\申万'
sw1_dir = r'D:\CC\Mid\估值\中信\SW1'
C_PRI=RGBColor(27,54,93);C_SEC=RGBColor(74,119,122);C_RED=RGBColor(192,57,43);C_GRN=RGBColor(39,174,96)

# ════════════════════════════════════════════════════════════
# 1. 读取PE/PB原始数据
# ════════════════════════════════════════════════════════════
def read_pe_pb(sheet_name, file_suffix):
    fp = os.path.join(r'D:\CC\DB\MKT', f'申万行业指数{file_suffix}_2026.xlsx')
    raw = pd.read_excel(fp, sheet_name=sheet_name, header=None)
    names = [str(raw.iloc[1, j]).split('(')[0] for j in range(1, raw.shape[1])]
    dates = pd.to_datetime(raw.iloc[2:, 0])
    data = raw.iloc[2:, 1:].copy().reset_index(drop=True)
    data.columns = names
    for c in names:
        data[c] = pd.to_numeric(data[c], errors='coerce')
    data.index = dates
    data = data[data.index.notna()]
    return names, data

pe_names, pe_data = read_pe_pb('PEtxt', 'PE')
pb_names, pb_data = read_pe_pb('PBtxt', 'PB')
print(f"PE数据: {pe_data.shape[0]}行 {len(pe_names)}行业 {pe_data.index.min().date()}~{pe_data.index.max().date()}")
print(f"PB数据: {pb_data.shape[0]}行 {len(pb_names)}行业 {pb_data.index.min().date()}~{pb_data.index.max().date()}")

# ════════════════════════════════════════════════════════════
# 2. 计算PE/PB统计值
# ════════════════════════════════════════════════════════════
def calc_stats(data, name):
    latest = data.iloc[-1]
    rows = []
    for c in data.columns:
        vals = data[c].dropna()
        if len(vals) < 2:
            continue
        pct = (vals.sort_values() < latest[c]).mean() * 100
        rows.append({
            '行业': c,
            f'最新{name}': round(latest[c], 2),
            f'历史最高{name}': round(vals.max(), 2),
            f'历史最低{name}': round(vals.min(), 2),
            f'{name}1/4分位': round(vals.quantile(0.25), 2),
            f'{name}中位数': round(vals.quantile(0.50), 2),
            f'{name}3/4分位': round(vals.quantile(0.75), 2),
            f'{name}标准差': round(vals.std(), 2),
            f'{name}历史分位数(%)': round(pct, 2),
        })
    return pd.DataFrame(rows)

pe_stats = calc_stats(pe_data, 'PE')
pb_stats = calc_stats(pb_data, 'PB')
print(f"PE统计: {len(pe_stats)}行业, PB统计: {len(pb_stats)}行业")

pe_stats.to_excel(os.path.join(out_dir, '申万行业PE统计.xlsx'), index=False)
pb_stats.to_excel(os.path.join(out_dir, '申万行业PB统计.xlsx'), index=False)

# ════════════════════════════════════════════════════════════
# 3. 读取收益率数据
# ════════════════════════════════════════════════════════════
yr = pd.read_excel(os.path.join(sw1_dir, '申万一级行业年度收益率矩阵.xlsx'), index_col=0)
qr = pd.read_excel(os.path.join(sw1_dir, '申万一级行业季度收益率矩阵.xlsx'), index_col=0)
# 分离行业名称和代码（名称用于merge，代码留存备用）
yri = yr.index.astype(str)
yr.index = yri.str.split('\n').str[0].str.strip()
yr_code = yri.str.split('\n').str[1].str.strip()  # 行业代码
qri = qr.index.astype(str)
qr.index = qri.str.split('\n').str[0].str.strip()
qr_code = qri.str.split('\n').str[1].str.strip()  # 行业代码

ny = yr.shape[1]
stats = pd.DataFrame({
    '年化平均收益率(%)': yr.mean(1).round(2),
    '年化波动率(%)': yr.std(1).round(2),
    '最大年涨幅(%)': yr.max(1).round(2),
    '最大年跌幅(%)': yr.min(1).round(2),
    '累计收益率(%)': ((yr/100+1).prod(1)-1).mul(100).round(2)})
stats['年胜率(%)'] = (yr > 0).sum(1) / ny * 100
stats['收益风险比'] = (stats['累计收益率(%)'] / (-stats['最大年跌幅(%)'].clip(upper=-0.01))).round(2)

latest_yr_cols = [c for c in yr.columns if c >= '2024']
stats['近3年累计收益(%)'] = ((yr[latest_yr_cols]/100+1).prod(1)-1).mul(100).round(2)
if '2026' in yr.columns: stats['2026年收益(%)'] = yr['2026'].round(2)
if '2025' in yr.columns: stats['2025年收益(%)'] = yr['2025'].round(2)
latest_q = qr.columns[-1] if len(qr.columns) > 0 else ''
if latest_q: stats['最新季度收益(%)'] = qr[latest_q].round(2)
if len(qr.columns) >= 4:
    stats['近4季累计收益(%)'] = ((qr[qr.columns[-4:]]/100+1).prod(1)-1).mul(100).round(2)

stats = stats.reset_index().rename(columns={'index': '行业'})

# ════════════════════════════════════════════════════════════
# 4. PE+PB综合评分
# ════════════════════════════════════════════════════════════
merged = stats.merge(pe_stats, on='行业', how='left').merge(pb_stats, on='行业', how='left')

merged['PE评分'] = pd.cut(merged['PE历史分位数(%)'].fillna(50), bins=[-1, 30, 50, 70, 100], labels=[3, 2, 1, 0]).astype(int)
merged['PB评分'] = pd.cut(merged['PB历史分位数(%)'].fillna(50), bins=[-1, 30, 50, 70, 100], labels=[3, 2, 1, 0]).astype(int)
merged['估值评分'] = merged['PE评分'] * 0.6 + merged['PB评分'] * 0.4

mom_cols = [c for c in ['近3年累计收益(%)', '2026年收益(%)', '最新季度收益(%)'] if c in merged.columns]
if mom_cols:
    merged['动量综合'] = merged[mom_cols].mean(1)
    merged['动量评分'] = pd.qcut(merged['动量综合'].rank(method='first'), q=4, labels=[0, 1, 2, 3]).astype(int)
else:
    merged['动量评分'] = 1; merged['动量综合'] = 0

merged['风险评分'] = (merged['年化波动率(%)'].rank(pct=True) * 3).round(0).clip(0, 3).astype(int)
merged['综合评分'] = (merged['估值评分']*0.4 + merged['动量评分']*0.4 + (3-merged['风险评分'])*0.2).round(2)

def classify(s):
    if s >= 2.0: return '积极配置'
    elif s >= 1.5: return '标配'
    elif s >= 1.0: return '谨慎观望'
    else: return '规避'
merged['投资评级'] = merged['综合评分'].apply(classify)

def pct_label(p, t='PE'):
    if p <= 20: return f'{t}深度低估'
    elif p <= 40: return f'{t}低估'
    elif p <= 60: return f'{t}合理'
    elif p <= 80: return f'{t}偏高'
    else: return f'{t}高估'

merged['PE状态'] = merged['PE历史分位数(%)'].fillna(50).apply(lambda p: pct_label(p, 'PE'))
merged['PB状态'] = merged['PB历史分位数(%)'].fillna(50).apply(lambda p: pct_label(p, 'PB'))

def mom_label(r):
    return {3: '强势', 2: '偏强', 1: '偏弱', 0: '弱势'}.get(r, '未知')
merged['动量状态'] = merged['动量评分'].apply(mom_label)
merged['风险等级'] = merged['风险评分'].apply(lambda r: {0: '低', 1: '低', 2: '中', 3: '高'}.get(r, '未知'))

merged = merged.sort_values('综合评分', ascending=False)

# ════════════════════════════════════════════════════════════
# 5. 导出Excel (PE+PB综合)
# ════════════════════════════════════════════════════════════
ex_cols = ['行业', '投资评级', '综合评分', '估值评分', '动量评分', '风险等级',
           '最新PE', 'PE历史分位数(%)', 'PE状态', '最新PB', 'PB历史分位数(%)', 'PB状态',
           '动量状态', '2026年收益(%)', '近3年累计收益(%)', '年化波动率(%)', '年胜率(%)', '收益风险比']
ex_cols = [c for c in ex_cols if c in merged.columns]
excel_path = os.path.join(out_dir, '行业投资价值综合评估_2026.xlsx')
with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
    merged[ex_cols].to_excel(writer, sheet_name='PE+PB综合', index=False)
print(f"导出: {excel_path}")

# ════════════════════════════════════════════════════════════
# 6. 导出Excel (PB-only)
# ════════════════════════════════════════════════════════════
merged2 = merged.copy()
merged2['估值评分'] = merged2['PB评分']
merged2['综合评分'] = (merged2['估值评分']*0.4 + merged2['动量评分']*0.4 + (3-merged2['风险评分'])*0.2).round(2)
merged2['投资评级'] = merged2['综合评分'].apply(classify)
merged2 = merged2.sort_values('综合评分', ascending=False)

ex_cols2 = ['行业', '投资评级', '综合评分', '估值评分', '动量评分', '风险等级',
            '最新PB', 'PB历史分位数(%)', 'PB状态',
            '动量状态', '2026年收益(%)', '近3年累计收益(%)', '年化波动率(%)', '年胜率(%)', '收益风险比']
ex_cols2 = [c for c in ex_cols2 if c in merged2.columns]
excel_path2 = os.path.join(out_dir, '行业投资价值综合评估_PB_2026.xlsx')
with pd.ExcelWriter(excel_path2, engine='openpyxl') as writer:
    merged2[ex_cols2].to_excel(writer, sheet_name='PB综合', index=False)
print(f"导出: {excel_path2}")

# ════════════════════════════════════════════════════════════
# 7. Word报告函数
# ════════════════════════════════════════════════════════════
def sf(run, name='Microsoft YaHei', size=Pt(10.5), color=RGBColor(51,51,51), bold=False, italic=False):
    run.font.name = name; run.font.size = size; run.font.color.rgb = color; run.bold = bold; run.italic = italic
    rPr = run._r.get_or_add_rPr(); rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}"/>'))

def add_title(doc, text, size=Pt(22)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); sf(r, size=size, color=C_PRI, bold=True)

def add_h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f'一、{text}'); sf(r, size=Pt(14), color=C_PRI, bold=True)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="1B365D"/></w:pBdr>'))

def add_h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'▎{text}'); sf(r, size=Pt(11), color=C_SEC, bold=True)

def add_body(doc, text, bf=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.7)
    if bf: r = p.add_run(bf); sf(r, bold=True, color=C_PRI)
    r = p.add_run(text); sf(r)

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def fmt_table(t):
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(8)
            if i == 0:
                shade(cell, '1B365D')
                for p in cell.paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(255,255,255); r.bold = True

def add_table(doc, df, title, cols):
    add_h2(doc, title)
    n = len(df) + 1; t = doc.add_table(rows=n, cols=len(cols)); t.style = 'Table Grid'
    for j, h in enumerate(cols): t.cell(0, j).text = h
    for i, (_, row) in enumerate(df.iterrows()):
        for j, c in enumerate(cols):
            v = row.get(c, '')
            t.cell(i+1, j).text = str(v) if not pd.isna(v) else '-'
            if c == '投资评级' and isinstance(v, str):
                clr = C_GRN if '积极' in v else (C_SEC if '标配' in v else (RGBColor(200,150,30) if '观望' in v else C_RED))
                for p in t.cell(i+1, j).paragraphs:
                    for r in p.runs: r.font.color.rgb = clr; r.bold = True
    fmt_table(t)

def add_quad(doc, m, t1, t2, label, cols):
    q = m[(m['估值评分'] >= t1) & (m['动量评分'] >= t2)]
    if len(q):
        add_h2(doc, label)
        add_table(doc, q[cols], f'{len(q)}个行业', cols)

def make_word_report(m, title, subtitle, tag, rcols, is_pb=False):
    doc = Document()
    add_title(doc, title)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); sf(r, size=Pt(11), color=C_SEC)
    doc.add_paragraph()

    add_h1(doc, '分析框架说明')
    if is_pb:
        add_body(doc, '估值维度仅使用PB（市净率）历史百分位数。')
    else:
        add_body(doc, '估值维度综合PE（60%）和PB（40%）历史百分位数。')
    add_body(doc, '动量维度基于2025-2026年收益、近3年累计收益。风险基于年化波动率。')
    add_body(doc, '综合评分=估值×40%+动量×40%+(3-风险)×20%。')

    add_h1(doc, '综合排名')
    add_table(doc, m.head(31), '全行业排名', rcols)

    for level, label in [('积极配置','积极配置板块'),('标配','标配板块'),('谨慎观望','谨慎观望板块'),('规避','规避板块')]:
        sub = m[m['投资评级'] == level]
        if len(sub): add_h1(doc, label); add_table(doc, sub[rcols], '', rcols)

    add_h1(doc, '分维度分析')
    if is_pb:
        add_table(doc, m.sort_values('PB历史分位数(%)')[['行业','PB状态','最新PB','PB历史分位数(%)']], 'PB估值由低到高', ['行业','PB状态','最新PB','PB历史分位数(%)'])
    else:
        add_table(doc, m.sort_values('PE历史分位数(%)')[['行业','PE状态','最新PE','PE历史分位数(%)','PB状态','最新PB','PB历史分位数(%)']], 'PE估值由低到高', ['行业','PE状态','最新PE','PE历史分位数(%)','PB状态','最新PB','PB历史分位数(%)'])

    mom_d = [c for c in ['行业','动量状态','2026年收益(%)','近3年累计收益(%)','最新季度收益(%)'] if c in m.columns]
    add_table(doc, m.sort_values('动量综合', ascending=False)[mom_d], '动量排名', mom_d)
    add_table(doc, m.sort_values('年化波动率(%)')[['行业','风险等级','年化波动率(%)','最大年跌幅(%)','年胜率(%)','收益风险比']], '风险排名', ['行业','风险等级','年化波动率(%)','最大年跌幅(%)','年胜率(%)','收益风险比'])

    add_h1(doc, '四象限分析')
    qcols = ['行业','投资评级','最新PE','PE历史分位数(%)','最新PB','PB历史分位数(%)','近3年累计收益(%)']
    qcols = [c for c in qcols if c in m.columns]
    add_quad(doc, m, 2, 2, '估值偏低+动量强势（超配）', qcols)
    add_quad(doc, m, 0, 2, '估值偏高+动量强势（趋势跟踪）', qcols)
    add_quad(doc, m, 2, 0, '估值偏低+动量弱势（左侧观察）', qcols)
    add_quad(doc, m, 0, 0, '估值偏高+动量弱势（规避）', qcols)

    add_h1(doc, '投资策略总结')
    active = m[m['投资评级']=='积极配置']['行业'].tolist()
    neutral = m[m['投资评级']=='标配']['行业'].tolist()
    av = m[m['投资评级'].isin(['谨慎观望','规避'])]['行业'].tolist()
    if active: add_body(doc, '、'.join(active)+'。', '积极配置（超配）：')
    if neutral: add_body(doc, '、'.join(neutral)+'。', '标配：')
    if av: add_body(doc, '、'.join(av)+'。', '谨慎/规避（低配）：')

    out = os.path.join(out_dir, f'申万一级行业综合投资价值与风险评估报告_{tag}_2026.docx')
    doc.save(out); print(f"导出: {out}")

# ════════════════════════════════════════════════════════════
# 8. 生成Word报告 (PE+PB综合版 & PB版)
# ════════════════════════════════════════════════════════════
r1 = ['行业','投资评级','综合评分','PE状态','PE历史分位数(%)','PB状态','PB历史分位数(%)','动量状态','风险等级']
make_word_report(merged, '申万一级行业综合投资价值与风险评估报告',
    f'基于最新PE/PB数据（截至{pe_data.index.max().date()}）+ 2016-2026收益率', 'PE_PB版', r1)

r2 = ['行业','投资评级','综合评分','PB状态','PB历史分位数(%)','动量状态','风险等级']
make_word_report(merged2, '申万一级行业综合投资价值与风险评估报告（PB版）',
    f'基于最新PB数据（截至{pb_data.index.max().date()}）+ 2016-2026收益率', 'PB版', r2, is_pb=True)

print("\n全部完成！")
