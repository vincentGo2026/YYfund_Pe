# -*- coding: utf-8 -*-
"""导出季度收益率分析报告 Word 文档"""
import pandas as pd, numpy as np, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

out_dir = r'D:\CC\Mid\估值\申万'
df = pd.read_excel(os.path.join(out_dir,'申万一级行业季度收益率矩阵.xlsx'),index_col=0)

s=pd.DataFrame({'平均收益率(%)':df.mean(1).round(2),'波动率(%)':df.std(1).round(2),
    '最大涨幅(%)':df.max(1).round(2),'最大跌幅(%)':df.min(1).round(2),
    '累计收益率(%)':((df/100+1).prod(1)-1).mul(100).round(2)})
s['胜率(%)']=(df>0).sum(1)/18*100
s['收益风险比']=(s['累计收益率(%)']/(-s['最大跌幅(%)'].clip(upper=-0.01))).round(2)
s=s.sort_values('累计收益率(%)',ascending=False)

doc=Document()
C_PRI=RGBColor(27,54,93);C_SEC=RGBColor(74,119,122);C_RED=RGBColor(192,57,43);C_GRN=RGBColor(39,174,96)

def sf(run,name='Microsoft YaHei',size=Pt(10.5),color=RGBColor(51,51,51),bold=False,italic=False):
    run.font.name=name;run.font.size=size;run.font.color.rgb=color;run.bold=bold;run.italic=italic
    r=run._r.get_or_add_rPr();r.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}"/>'))

def add_t(text,size=Pt(22)):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text);sf(r,size=size,color=C_PRI,bold=True)

def h1(text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(16);p.paragraph_format.space_after=Pt(8)
    r=p.add_run(f'一、{text}');sf(r,size=Pt(14),color=C_PRI,bold=True)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="1B365D"/></w:pBdr>'))

def h2(text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f'▎{text}');sf(r,size=Pt(11),color=C_SEC,bold=True)

def body(text,bf=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(4);p.paragraph_format.first_line_indent=Cm(0.7)
    if bf:r=p.add_run(bf);sf(r,bold=True,color=C_PRI)
    r=p.add_run(text);sf(r)

def shade(cell,color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def ft(t):
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate(t.rows):
        for j,cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:r.font.size=Pt(8.5)
            if i==0:
                shade(cell,'1B365D')
                for p in cell.paragraphs:
                    for r in p.runs:r.font.color.rgb=RGBColor(255,255,255);r.bold=True

def tbl(data,title):
    h2(title);n=len(data)+1;t=doc.add_table(rows=n,cols=6);t.style='Table Grid'
    hs=['排名','行业','累计收益率','平均季收益','波动率','胜率']
    for j,h in enumerate(hs):t.cell(0,j).text=h
    for i,(idx,row) in enumerate(data.iterrows()):
        t.cell(i+1,0).text=str(i+1);t.cell(i+1,1).text=str(idx)
        t.cell(i+1,2).text=f"{row['累计收益率(%)']:.1f}%"
        t.cell(i+1,3).text=f"{row['平均收益率(%)']:.2f}%"
        t.cell(i+1,4).text=f"{row['波动率(%)']:.1f}%"
        t.cell(i+1,5).text=f"{row['胜率(%)']:.0f}%"
        c=C_GRN if row['累计收益率(%)']>0 else C_RED
        for j in[2,3]:
            for p in t.cell(i+1,j).paragraphs:
                for r in p.runs:r.font.color.rgb=c;r.bold=True
    ft(t)

add_t('申万一级行业季度收益率投资价值与风险评估报告')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('基于2022Q1-2026Q2季度收益率热力图分析');sf(r,size=Pt(11),color=C_SEC)
doc.add_paragraph()

h1('行业综合排名')
body('以下基于2022Q1至2026Q2共18个季度数据，对31个申万一级行业进行收益率统计分析。排名按累计收益率降序排列。')
tbl(s,'全部31个行业排名')

h1('行业投资价值分类')
h2('🟢 第一梯队：进攻型（高收益高波动）')
top3=s.head(6);tbl(top3,'Top 6行业')
body('通信以+236.3%累计收益率绝对领跑，驱动来自AI算力基础设施、光通信、5G/6G产业链持续爆发。银行（+41.7%）和煤炭（+28.1%）代表了高股息+低估值修复逻辑。有色金属（+18.7%）受益大宗商品牛市，机械设备（+23.9%）受益制造业升级和出海。')

h2('🟡 第二梯队：稳健型（低波动绝对收益）')
mid=s.iloc[6:15];tbl(mid,'第7-15名')
body('公用事业波动率仅6.8%全市场最低，具备类债券属性。石油石化、家用电器累计收益率接近零但胜率较高（50%+），适合稳健底仓。')

h2('🔴 第三梯队：困境型（大幅亏损等待拐点）')
bot=s.tail(8);tbl(bot,'排名靠后行业')
body('房地产胜率仅17%（18季中仅3季正收益），行业逻辑未理顺。医药生物(-46.3%)、食品饮料(-46.6%)、美容护理(-59.3%)受消费降级与估值回归双重压制。')

h1('风险评估矩阵')
h2('高波动行业（季度波动率>12%）')
tbl(s[s['波动率(%)']>12].sort_values('波动率(%)',ascending=False),'高波动行业')
h2('高回撤行业（最大单季跌幅>20%）')
tbl(s[s['最大跌幅(%)']<-20].sort_values('最大跌幅(%)'),'高回撤行业')
h2('收益风险比排名')
tbl(s.sort_values('收益风险比',ascending=False),'性价比排名')
body('通信以20.2的收益风险比遥遥领先，每承担1%回撤可获得20.2%累计收益。')

h1('投资策略建议')
body('通信逢回调布局，有色金属和煤炭可作为抗通胀卫星配置。','进攻型配置：')
body('公用事业作为债券替代，家用电器胜率56%，石油石化高股息。','核心底仓配置：')
body('食品饮料和医药已连续下跌3年+估值低位，但需消费数据验证拐点。','困境反转观察：')
body('美容护理未见止跌迹象，房地产行业逻辑未理顺，建议规避。','需要规避：')

h1('季度收益率热力图')
img=os.path.join(out_dir,'申万一级行业季度收益率热力图.png')
if os.path.exists(img):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img,width=Inches(5.5))
    p2=doc.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p2.add_run('图：申万一级行业季度收益率热力图 (2022Q1-2026Q2)');sf(r,size=Pt(9),color=C_SEC,italic=True)

out=os.path.join(out_dir,'申万一级行业季度收益率分析报告.docx')
doc.save(out);print(f"导出: {out}")
